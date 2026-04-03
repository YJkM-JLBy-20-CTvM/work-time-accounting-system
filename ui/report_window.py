import sys
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QMessageBox,
    QComboBox, QDateEdit, QFileDialog
)
from PyQt5.QtCore import QDate
from db_connection import get_connection
from docx import Document
from datetime import datetime


class ReportWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Формирование отчётов")
        self.resize(400, 300)

        layout = QVBoxLayout()

        self.report_type = QComboBox()
        self.report_type.addItems([
            "По сотруднику",
            "По отделу"
        ])

        self.employee_combo = QComboBox()
        layout.addWidget(QLabel("Сотрудник:"))
        layout.addWidget(self.employee_combo)

        self.department_combo = QComboBox()
        layout.addWidget(QLabel("Отдел:"))
        layout.addWidget(self.department_combo)

        date_layout = QHBoxLayout()

        self.date_from = QDateEdit()
        self.date_from.setCalendarPopup(True)
        self.date_from.setDate(QDate.currentDate().addMonths(-1))

        self.date_to = QDateEdit()
        self.date_to.setCalendarPopup(True)
        self.date_to.setDate(QDate.currentDate())

        date_layout.addWidget(QLabel("С:"))
        date_layout.addWidget(self.date_from)
        date_layout.addWidget(QLabel("По:"))
        date_layout.addWidget(self.date_to)

        layout.addLayout(date_layout)

        self.generate_btn = QPushButton("Сформировать отчёт")
        layout.addWidget(self.generate_btn)

        self.setLayout(layout)

        self.generate_btn.clicked.connect(self.generate_report)

        self.load_employees()
        self.load_departments()

    def load_employees(self):
        conn = get_connection()
        cur = conn.cursor()

        try:
            cur.execute("""
                SELECT employee_id, last_name, first_name
                FROM employees
                ORDER BY last_name
            """)

            self.employee_combo.addItem("Все сотрудники", None)

            for emp_id, last, first in cur.fetchall():
                self.employee_combo.addItem(f"{last} {first}", emp_id)

        finally:
            cur.close()
            conn.close()

    def load_departments(self):
        conn = get_connection()
        cur = conn.cursor()

        try:
            cur.execute("""
                SELECT department_id, department_name
                FROM departments
                ORDER BY department_name
            """)

            self.department_combo.addItem("Все отделы", None)

            for dep_id, name in cur.fetchall():
                self.department_combo.addItem(name, dep_id)

        finally:
            cur.close()
            conn.close()

    def generate_report(self):
        conn = get_connection()
        cur = conn.cursor()

        try:
            employee_id = self.employee_combo.currentData()
            department_id = self.department_combo.currentData()

            date_from = self.date_from.date().toPyDate()
            date_to = self.date_to.date().toPyDate()

            query = """
                SELECT e.employee_id,
                    e.last_name, e.first_name,
                    d.department_name,
                    w.record_date,
                    w.arrival_time,
                    w.departure_time
                FROM employees e
                JOIN departments d ON e.department_id = d.department_id
                JOIN work_time_accounting w ON e.employee_id = w.employee_id
                WHERE w.record_date BETWEEN %s AND %s
            """

            params = [date_from, date_to]

            if employee_id:
                query += " AND e.employee_id = %s"
                params.append(employee_id)

            if department_id:
                query += " AND e.department_id = %s"
                params.append(department_id)

            query += " ORDER BY e.last_name, w.record_date"

            cur.execute(query, params)
            rows = cur.fetchall()

            if not rows:
                QMessageBox.warning(self, "Отчёт", "Нет данных")
                return

            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Сохранить отчёт",
                "",
                "Документ Word (*.docx)"
            )

            if not file_path:
                return

            if not file_path.endswith(".docx"):
                file_path += ".docx"

            doc = Document()

            title = doc.add_heading("Отчёт по рабочему времени", 0)
            title.alignment = 1

            doc.add_paragraph(
                f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            )
            doc.add_paragraph("")

            headers = [
                "Фамилия", "Имя", "Отдел",
                "Дата", "Приход", "Уход", "Часы"
            ]

            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = "Table Grid"

            for col, header in enumerate(headers):
                table.cell(0, col).text = header

            employee_totals = {}

            total_minutes_all = 0

            for i, row in enumerate(rows):
                emp_id, last_name, first_name, department, date, arrival, departure = row

                if arrival and departure:
                    total_minutes = (
                        departure.hour * 60 + departure.minute
                    ) - (
                        arrival.hour * 60 + arrival.minute
                    )

                    hours = total_minutes // 60
                    minutes = total_minutes % 60

                    work_time = f"{hours} ч {minutes} мин"

                    total_minutes_all += total_minutes

                    if emp_id not in employee_totals:
                        employee_totals[emp_id] = 0
                    employee_totals[emp_id] += total_minutes

                else:
                    work_time = "—"

                values = [
                    last_name,
                    first_name,
                    department,
                    date.strftime("%d.%m.%Y"),
                    arrival.strftime("%H:%M") if arrival else "—",
                    departure.strftime("%H:%M") if departure else "—",
                    work_time
                ]

                for j, value in enumerate(values):
                    table.cell(i + 1, j).text = str(value)

            doc.add_paragraph("\nИтого по сотрудникам:")

            for emp_id, minutes in employee_totals.items():
                hours = minutes // 60
                mins = minutes % 60

                cur.execute("""
                    SELECT last_name, first_name
                    FROM employees
                    WHERE employee_id = %s
                """, (emp_id,))
                last, first = cur.fetchone()

                doc.add_paragraph(f"{last} {first}: {hours} ч {mins} мин")

            total_hours = total_minutes_all // 60
            total_minutes = total_minutes_all % 60

            doc.add_paragraph("\n----------------------")
            doc.add_paragraph(f"ОБЩИЙ ИТОГ: {total_hours} ч {total_minutes} мин")

            doc.save(file_path)

            QMessageBox.information(self, "Готово", "Отчёт создан")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

        finally:
            cur.close()
            conn.close()