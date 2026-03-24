from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QLabel, QPushButton,
    QComboBox, QTextEdit, QMessageBox, QFileDialog
)
from docx import Document
from datetime import datetime
from db_connection import get_connection
from ui.employee_crud_window import EmployeeCRUD
from ui.department_crud_window import DepartmentCRUD
from ui.position_crud_window import PositionCRUD


class AdminWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Окно администратора")
        self.resize(500, 500)

        layout = QVBoxLayout()

        layout.addWidget(QLabel("Выберите сотрудника:"))

        self.employee_combo = QComboBox()
        layout.addWidget(self.employee_combo)

        self.load_employees()

        self.btn_view = QPushButton("Посмотреть рабочее время")
        self.btn_view.clicked.connect(self.view_employee_time)
        layout.addWidget(self.btn_view)

        self.output = QTextEdit()
        self.output.setReadOnly(True)
        layout.addWidget(self.output)

        self.btn_crud = QPushButton("Открыть меню редактирования таблиц")
        self.btn_crud.clicked.connect(self.open_crud)
        layout.addWidget(self.btn_crud)

        layout.addWidget(QLabel("Сформировать отчёт:"))

        self.report_type = QComboBox()
        self.report_type.addItems(["По сотруднику", "По отделу"])
        layout.addWidget(self.report_type)

        self.btn_report = QPushButton("Сформировать")
        self.btn_report.clicked.connect(self.generate_report)
        layout.addWidget(self.btn_report)

        self.setLayout(layout)

    def load_employees(self):
        conn = get_connection()
        cur = conn.cursor()

        try:
            cur.execute("""
                SELECT employee_id, last_name, first_name
                FROM employees
            """)

            self.employees = cur.fetchall()

            for emp in self.employees:
                emp_id, last, first = emp
                self.employee_combo.addItem(f"{last} {first}", emp_id)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

        finally:
            cur.close()
            conn.close()

    def view_employee_time(self):
        employee_id = self.employee_combo.currentData()

        conn = get_connection()
        cur = conn.cursor()

        try:
            cur.execute("""
                SELECT record_date, arrival_time, departure_time
                FROM work_time_accounting
                WHERE employee_id = %s
                ORDER BY record_date DESC
            """, (employee_id,))

            records = cur.fetchall()

            self.output.clear()

            for record in records:
                date, arrival, departure = record

                arrival_str = arrival.strftime("%H:%M") if arrival else "---"
                departure_str = departure.strftime("%H:%M") if departure else "---"

                self.output.append(
                    f"{date} | {arrival_str} - {departure_str}"
                )

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

        finally:
            cur.close()
            conn.close()

    def open_crud(self):
        self.crud_menu = QWidget()
        self.crud_menu.setWindowTitle("Выбор таблицы")
        
        layout = QVBoxLayout()
        
        btn_emp = QPushButton("Сотрудники")
        btn_dep = QPushButton("Отделы")
        btn_pos = QPushButton("Должности")
        
        btn_emp.clicked.connect(self.open_employee_crud)
        btn_dep.clicked.connect(self.open_department_crud)
        btn_pos.clicked.connect(self.open_position_crud)
        
        layout.addWidget(btn_emp)
        layout.addWidget(btn_dep)
        layout.addWidget(btn_pos)
        
        self.crud_menu.setLayout(layout)
        self.crud_menu.show()
        
    def open_employee_crud(self):
        self.employee_crud = EmployeeCRUD()
        self.employee_crud.show()
        
    def open_department_crud(self):
        self.department_crud = DepartmentCRUD()
        self.department_crud.show()
        
    def open_position_crud(self):
        self.position_crud = PositionCRUD()
        self.position_crud.show()

    def generate_report(self):
        conn = get_connection()
        cur = conn.cursor()

        try:
            report_type = self.report_type.currentText()

            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Сохранить отчёт",
                "",
                "Документ Word (*.docx)"
            )

            if not file_path:
                return

            if not file_path.endswith(".docx"):
                file_path += ".docx"

            doc = Document()

            title = doc.add_heading("Отчёт по рабочему времени", 0)
            title.alignment = 1

            doc.add_paragraph(
                f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            )
            doc.add_paragraph("")

            if report_type == "По сотруднику":
                employee_id = self.employee_combo.currentData()

                cur.execute("""
                    SELECT e.last_name, e.first_name,
                        w.record_date, w.arrival_time, w.departure_time
                    FROM employees e
                    JOIN work_time_accounting w
                        ON e.employee_id = w.employee_id
                    WHERE e.employee_id = %s
                    ORDER BY w.record_date
                """, (employee_id,))

                rows = cur.fetchall()

                headers = ["Фамилия", "Имя", "Дата", "Приход", "Уход"]

            else:
                department_id = self.department_combo.currentData()

                cur.execute("""
                    SELECT e.last_name, e.first_name,
                        w.record_date, w.arrival_time, w.departure_time
                    FROM employees e
                    JOIN work_time_accounting w
                        ON e.employee_id = w.employee_id
                    WHERE e.department_id = %s
                    ORDER BY e.last_name, w.record_date
                """, (department_id,))

                rows = cur.fetchall()

                headers = ["Фамилия", "Имя", "Дата", "Приход", "Уход"]

            if not rows:
                QMessageBox.warning(self, "Отчёт", "Нет данных для отчёта")
                return

            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = "Table Grid"

            for col, header in enumerate(headers):
                table.cell(0, col).text = header

            for i, row in enumerate(rows):
                for j, value in enumerate(row):
                    if hasattr(value, "strftime"):
                        if "time" in str(type(value)):
                            value = value.strftime("%H:%M")
                        else:
                            value = value.strftime("%d.%m.%Y")

                    table.cell(i + 1, j).text = str(value)

            doc.save(file_path)

            QMessageBox.information(self, "Готово", "Отчёт успешно создан")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

        finally:
            cur.close()
            conn.close()